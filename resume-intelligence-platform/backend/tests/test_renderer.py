import pytest
import io
import docx
from app.models.schema import ResumeDocument, ExperienceItem, ProjectItem, EducationItem
from app.services.renderer import render

def test_render_txt_serialization():
    doc = ResumeDocument(
        summary="Test summary.",
        skills=["Python", "FastAPI"],
        experience=[
            ExperienceItem(company="CompA", title="SWE", start_date="2020-01", end_date="Present", bullets=["Bullet 1"])
        ],
        projects=[
            ProjectItem(name="ProjA", description="Desc", tech_stack=["Python"], bullets=["B1"])
        ],
        education=[
            EducationItem(institution="MIT", degree="CS", start_date="2016", end_date="2020")
        ],
        certifications=["AWS Solutions Architect"]
    )
    
    txt_bytes = render(b"", "resume.docx", doc, "txt")
    text = txt_bytes.decode("utf-8")
    
    assert "SUMMARY" in text
    assert "Test summary." in text
    assert "TECHNICAL SKILLS" in text
    assert "Python, FastAPI" in text
    assert "CompA" in text
    assert "ProjA" in text
    assert "MIT" in text
    assert "AWS Solutions Architect" in text

def test_render_pdf_reportlab_generation():
    doc = ResumeDocument(
        summary="Professional engineer summary.",
        skills=["AWS", "Docker"],
        experience=[
            ExperienceItem(company="Cloudy Inc", title="SWE", start_date="2020", bullets=["Created lambda services."])
        ]
    )
    
    # Render PDF bytes
    pdf_bytes = render(b"", "resume.pdf", doc, "pdf")
    
    # Assert header bytes representing a valid PDF file
    assert pdf_bytes.startswith(b"%PDF-")

def test_render_docx_text_run_replacements():
    # 1. Create a mock docx in memory
    doc = docx.Document()
    doc.add_paragraph("Original Summary Text Block that represents candidate profile.")
    doc.add_paragraph("• Worked on some legacy systems.")
    
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    original_docx_bytes = docx_buf.getvalue()
    
    # 2. Setup customized document
    rewritten = ResumeDocument(
        summary="Modern Customized Summary showing python focus.",
        skills=[],
        experience=[
            ExperienceItem(company="A", title="SWE", start_date="2020", bullets=["Optimized Python web servers."])
        ]
    )
    
    rendered_docx_bytes = render(original_docx_bytes, "resume.docx", rewritten, "docx")
    
    # Open modified docx and assert text replacements
    modified_doc = docx.Document(io.BytesIO(rendered_docx_bytes))
    all_paras = [p.text for p in modified_doc.paragraphs]
    
    # Verify summary is modified
    assert "Modern Customized Summary showing python focus." in all_paras
    # Verify bullet is modified
    assert any("Optimized Python web servers." in p for p in all_paras)

def test_render_pdf_header_and_divider():
    doc = ResumeDocument(
        name="Rajat Singh",
        contact_info=["rajat@example.com", "1234567890", "github.com/rajat"],
        summary="Professional summary.",
        skills=["Python"]
    )
    pdf_bytes = render(b"", "resume.pdf", doc, "pdf")
    assert pdf_bytes.startswith(b"%PDF-")


def test_render_docx_table_cells_preservation():
    # Create docx with a table cell containing bullet points
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = "Summary paragraph in table."
    cell.add_paragraph("• Original Bullet One")
    
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    original_bytes = docx_buf.getvalue()
    
    original_doc = ResumeDocument(
        summary="Summary paragraph in table.",
        experience=[
            ExperienceItem(company="CompA", title="SWE", start_date="2020", bullets=["Original Bullet One"])
        ]
    )
    
    rewritten = ResumeDocument(
        summary="Rewritten Summary in table.",
        experience=[
            ExperienceItem(company="CompA", title="SWE", start_date="2020", bullets=["Rewritten Bullet One"])
        ]
    )
    
    output_docx = render(original_bytes, "resume.docx", rewritten, "docx", original_doc)
    modified_doc = docx.Document(io.BytesIO(output_docx))
    all_texts = [p.text for p in modified_doc.tables[0].cell(0, 0).paragraphs]
    
    assert "Rewritten Summary in table." in all_texts
    assert any("Rewritten Bullet One" in t for t in all_texts)

def test_render_docx_dynamic_bullet_count_changes():
    # Test insertion and deletion
    doc = docx.Document()
    doc.add_paragraph("Summary text")
    doc.add_paragraph("• Bullet A")
    doc.add_paragraph("• Bullet B")
    
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    original_bytes = docx_buf.getvalue()
    
    original_doc = ResumeDocument(
        summary="Summary text",
        experience=[
            ExperienceItem(company="A", title="SWE", start_date="2020", bullets=["Bullet A", "Bullet B"])
        ]
    )
    
    # 1. Test bullet addition (A and B become X, Y, Z)
    rewritten_add = ResumeDocument(
        summary="Summary text",
        experience=[
            ExperienceItem(company="A", title="SWE", start_date="2020", bullets=["Bullet X", "Bullet Y", "Bullet Z"])
        ]
    )
    output_add = render(original_bytes, "resume.docx", rewritten_add, "docx", original_doc)
    doc_add = docx.Document(io.BytesIO(output_add))
    texts_add = [p.text for p in doc_add.paragraphs]
    # Check that we have 3 bullets now
    assert any("Bullet X" in t for t in texts_add)
    assert any("Bullet Y" in t for t in texts_add)
    assert any("Bullet Z" in t for t in texts_add)
    assert len([t for t in texts_add if "Bullet" in t]) == 3
    
    # 2. Test bullet subtraction (A and B become only X)
    rewritten_sub = ResumeDocument(
        summary="Summary text",
        experience=[
            ExperienceItem(company="A", title="SWE", start_date="2020", bullets=["Bullet X"])
        ]
    )
    output_sub = render(original_bytes, "resume.docx", rewritten_sub, "docx", original_doc)
    doc_sub = docx.Document(io.BytesIO(output_sub))
    texts_sub = [p.text for p in doc_sub.paragraphs]
    assert any("Bullet X" in t for t in texts_sub)
    assert len([t for t in texts_sub if "Bullet" in t]) == 1

def test_render_docx_unicode_and_smart_quotes():
    doc = docx.Document()
    doc.add_paragraph("“Smart quotes summary.”")
    doc.add_paragraph("• Dealt with ﬁnancial charts.") # ligature fi
    
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    original_bytes = docx_buf.getvalue()
    
    # original doc has standard quotes and words
    original_doc = ResumeDocument(
        summary='"Smart quotes summary."',
        experience=[
            ExperienceItem(company="A", title="SWE", start_date="2020", bullets=["Dealt with financial charts."])
        ]
    )
    
    rewritten = ResumeDocument(
        summary="New summary.",
        experience=[
            ExperienceItem(company="A", title="SWE", start_date="2020", bullets=["Rewritten financial bullet."])
        ]
    )
    
    output = render(original_bytes, "resume.docx", rewritten, "docx", original_doc)
    doc_out = docx.Document(io.BytesIO(output))
    texts = [p.text for p in doc_out.paragraphs]
    
    assert "New summary." in texts
    assert any("Rewritten financial bullet." in t for t in texts)


def test_replace_paragraph_text_preserving_prefix():
    from app.services.renderer import replace_paragraph_text_preserving_prefix
    
    # Mock a paragraph with multiple runs (e.g. ["•", "\t", "Original text"])
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("•")
    p.add_run("\t")
    p.add_run("Original text")
    
    assert len(p.runs) == 3
    assert p.runs[0].text == "•"
    assert p.runs[1].text == "\t"
    assert p.runs[2].text == "Original text"
    
    replace_paragraph_text_preserving_prefix(p, "•\tNew customized text")
    
    # Assert runs structure and values
    assert len(p.runs) == 3
    assert p.runs[0].text == "•"
    assert p.runs[1].text == "\t"
    assert p.runs[2].text == "New customized text"
    
    # Also test handling when new spacing is requested but original prefix had no space
    p2 = doc.add_paragraph()
    p2.add_run("•")
    p2.add_run("Original text")
    
    assert len(p2.runs) == 2
    replace_paragraph_text_preserving_prefix(p2, "• New customized text")
    
    # First content run should now contain the space and new text
    assert p2.runs[0].text == "•"
    assert p2.runs[1].text == " New customized text"


