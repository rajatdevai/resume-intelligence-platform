import io
import os
import re
import pdfplumber
import fitz  # PyMuPDF
import docx
from typing import List, Optional, Dict, Any
from app.core.exceptions import ParserError
from app.models.schema import ResumeDocument, ExperienceItem, ProjectItem, EducationItem

HEADINGS_MAP = {
    "summary": ["summary", "professional summary", "objective", "career objective", "about me", "professional objective", "career summary"],
    "experience": ["experience", "employment", "work history", "professional experience", "work experience", "employment history"],
    "skills": ["skills", "technical skills", "tech stack", "core skills", "skills & expertise", "skills and expertise", "expertise"],
    "projects": ["projects", "personal projects", "academic work", "academic projects", "key projects", "featured projects"],
    "education": ["education", "academics", "academic history", "education background"],
    "certifications": ["certifications", "licenses", "certificates", "licenses & certifications", "licenses and certifications"]
}

TITLE_KEYWORDS = {
    "engineer", "developer", "manager", "analyst", "intern", "specialist", 
    "lead", "head", "director", "consultant", "architect", "programmer", "designer", "officer"
}

def parse_resume(file_bytes: bytes, filename: str) -> ResumeDocument:
    """
    Parses PDF and DOCX files to extract and structure their text content.
    """
    if not file_bytes or len(file_bytes) == 0:
        raise ParserError("Empty file")

    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        raw_text = _parse_pdf(file_bytes)
    elif ext == ".docx":
        raw_text = _parse_docx(file_bytes)
    else:
        raise ParserError("Unsupported file type")

    if not raw_text.strip():
        raise ParserError("No text could be extracted from the file")

    # Split into sections and structure
    doc = _structure_resume(raw_text)
    return doc

def _parse_pdf(file_bytes: bytes) -> str:
    # 1. Try with pdfplumber
    try:
        pdf_stream = io.BytesIO(file_bytes)
        with pdfplumber.open(pdf_stream) as pdf:
            if len(pdf.pages) == 0:
                raise ParserError("Empty PDF file")
                
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                return text
    except ParserError:
        raise
    except Exception as e:
        if not isinstance(e, AttributeError):
            err_msg = str(e).lower()
            if "encrypted" in err_msg or "password" in err_msg or "authenticate" in err_msg:
                raise ParserError("Password-protected PDF")
        pass

    # 2. Fallback to PyMuPDF (fitz)
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if doc.is_encrypted:
            raise ParserError("Password-protected PDF")
        
        text = ""
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        
        if text.strip():
            return text
    except ParserError:
        raise
    except Exception as e:
        err_msg = str(e).lower()
        if "encrypted" in err_msg or "password" in err_msg or "authenticate" in err_msg:
            raise ParserError("Password-protected PDF")
        raise ParserError(f"Corrupted or invalid PDF: {str(e)}")

    raise ParserError("No text could be extracted from the PDF")

def _parse_docx(file_bytes: bytes) -> str:
    try:
        doc_stream = io.BytesIO(file_bytes)
        doc = docx.Document(doc_stream)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(cell.text)
        return "\n".join(text)
    except Exception as e:
        raise ParserError(f"Corrupted or invalid DOCX: {str(e)}")

def _clean_heading(line: str) -> str:
    line = line.strip()
    # Remove trailing decorative elements, colons, vertical bars, or spaces
    line = re.sub(r'[\_\-\–\—\s\|:]+$', '', line)
    # Remove markdown headers (#)
    line = re.sub(r'^#+\s*', '', line)
    # Remove leading numbers like 1. or 2.1.
    line = re.sub(r'^\d+(\.\d+)*[\s\.\-:]+', '', line)
    # Remove leading bullets/symbols
    line = re.sub(r'^[•\-\*\d\.\s]+', '', line)
    return line.strip().lower()

def _match_heading_key(cleaned_line: str) -> Optional[str]:
    for key, synonyms in HEADINGS_MAP.items():
        if cleaned_line in synonyms:
            return key
    return None

def _extract_date_range(line: str) -> tuple[Optional[str], Optional[str], str]:
    # Match patterns like: "Jan 2020 - Present", "2020-01 - 2022-02", "2020 - 2022", "March 2018 to Present"
    date_pattern = r'\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4}-\d{2}|\d{4})\s*[\-\–\—to/]+\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}|\d{4}-\d{2}|\d{4}|Present|Current)\b'
    match = re.search(date_pattern, line, re.IGNORECASE)
    if match:
        start, end = match.group(1), match.group(2)
        cleaned = line[:match.start()] + " " + line[match.end():]
        return start.strip(), end.strip(), cleaned.strip()
    return None, None, line

def _split_title_company(text: str) -> tuple[str, str]:
    text = re.sub(r'\s*[\|\,\-\–\—]\s*', ' | ', text)
    parts = [p.strip() for p in text.split('|') if p.strip()]
    if not parts:
        parts = [p.strip() for p in text.split(' at ') if p.strip()]
    
    if len(parts) >= 2:
        part1, part2 = parts[0], parts[1]
        p1_words = set(part1.lower().split())
        p2_words = set(part2.lower().split())
        
        if p1_words.intersection(TITLE_KEYWORDS):
            return part2, part1  # company, title
        elif p2_words.intersection(TITLE_KEYWORDS):
            return part1, part2  # company, title
        return part1, part2
    elif len(parts) == 1:
        return parts[0], ""
    return "", ""

def _parse_skills(skills_lines: List[str]) -> List[str]:
    skills = []
    for line in skills_lines:
        parts = re.split(r'[,;\|•\-\*]+', line)
        for p in parts:
            p_clean = p.strip()
            if p_clean and len(p_clean) < 50:
                skills.append(p_clean)
    return skills

def _parse_experience(exp_lines: List[str]) -> List[ExperienceItem]:
    items = []
    current_item = None
    
    for line in exp_lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        is_bullet = line_str.startswith(('-', '•', '*', '+', 'o '))
        if is_bullet:
            bullet_text = re.sub(r'^[\-•\*+o\s]+', '', line_str).strip()
            if current_item:
                current_item["bullets"].append(bullet_text)
            else:
                current_item = {
                    "company": "Unknown Company",
                    "title": "Position",
                    "start_date": "Unknown",
                    "end_date": None,
                    "bullets": [bullet_text]
                }
        else:
            start, end, cleaned = _extract_date_range(line_str)
            if start:
                if current_item:
                    items.append(ExperienceItem(**current_item))
                
                company, title = _split_title_company(cleaned)
                current_item = {
                    "company": company or "Unknown Company",
                    "title": title or "Position",
                    "start_date": start,
                    "end_date": end,
                    "bullets": []
                }
            else:
                if current_item:
                    if current_item["bullets"]:
                        # This line is a continuation of the previous bullet point
                        current_item["bullets"][-1] += " " + line_str
                    else:
                        if current_item["company"] == "Unknown Company" or not current_item["company"]:
                            current_item["company"] = line_str
                        elif current_item["title"] == "Position" or not current_item["title"]:
                            current_item["title"] = line_str
                else:
                    company, title = _split_title_company(line_str)
                    current_item = {
                        "company": company or "Unknown Company",
                        "title": title or "Position",
                        "start_date": "Unknown",
                        "end_date": None,
                        "bullets": []
                    }
                    
    if current_item:
        items.append(ExperienceItem(**current_item))
        
    return items

def _parse_projects(proj_lines: List[str]) -> List[ProjectItem]:
    items = []
    current_item = None
    
    for line in proj_lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        is_bullet = line_str.startswith(('-', '•', '*', '+', 'o '))
        if is_bullet:
            bullet_text = re.sub(r'^[\-•\*+o\s]+', '', line_str).strip()
            if current_item:
                current_item["bullets"].append(bullet_text)
            else:
                current_item = {
                    "name": "Project",
                    "description": "Project details",
                    "tech_stack": [],
                    "bullets": [bullet_text]
                }
        else:
            if current_item and current_item["bullets"]:
                # This line is a continuation of the previous bullet point
                current_item["bullets"][-1] += " " + line_str
            else:
                if current_item:
                    items.append(ProjectItem(**current_item))
                    
                tech_stack = []
                tech_match = re.search(r'\((.*?)\)', line_str)
                name = line_str
                if tech_match:
                    tech_str = tech_match.group(1)
                    tech_stack = [t.strip() for t in re.split(r'[,/|]+', tech_str) if t.strip()]
                    name = line_str[:tech_match.start()].strip()
                    
                current_item = {
                    "name": name or "Project",
                    "description": name or "Project description",
                    "tech_stack": tech_stack,
                    "bullets": []
                }
            
    if current_item:
        items.append(ProjectItem(**current_item))
        
    return items

def _parse_education(edu_lines: List[str]) -> List[EducationItem]:
    items = []
    for line in edu_lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        start, end, cleaned = _extract_date_range(line_str)
        parts = [p.strip() for p in re.split(r'[,||\-–\—]+', cleaned) if p.strip()]
        
        institution = "University"
        degree = "Degree"
        if len(parts) >= 2:
            institution = parts[0]
            degree = parts[1]
        elif len(parts) == 1:
            institution = parts[0]
            
        items.append(EducationItem(
            institution=institution,
            degree=degree,
            start_date=start or "Unknown",
            end_date=end
        ))
    return items

def _parse_certifications(cert_lines: List[str]) -> List[str]:
    certs = []
    for line in cert_lines:
        line_str = line.strip()
        if not line_str:
            continue
        clean_line = re.sub(r'^[\-•\*+o\s]+', '', line_str).strip()
        if clean_line:
            certs.append(clean_line)
    return certs

def _structure_resume(raw_text: str) -> ResumeDocument:
    lines = raw_text.split('\n')
    sections = {
        "summary": [],
        "skills": [],
        "experience": [],
        "projects": [],
        "education": [],
        "certifications": []
    }
    
    pre_header_lines = []
    current_section = None
    
    for line in lines:
        cleaned = _clean_heading(line)
        matched_key = _match_heading_key(cleaned)
        
        if matched_key:
            current_section = matched_key
        elif current_section:
            sections[current_section].append(line)
        else:
            if line.strip():
                pre_header_lines.append(line)

    # Fallback summary if no summary section heading is found
    summary_text = ""
    if sections["summary"]:
        summary_text = "\n".join(sections["summary"]).strip()
    elif pre_header_lines:
        # Use first 3 lines of headerless text as fallback summary if they aren't parsed into other things
        summary_text = " ".join([l.strip() for l in pre_header_lines[:4]])
        
    skills = _parse_skills(sections["skills"])
    experience = _parse_experience(sections["experience"])
    projects = _parse_projects(sections["projects"])
    education = _parse_education(sections["education"])
    certifications = _parse_certifications(sections["certifications"])
    
    parse_warnings = []
    
    # Confident parsing warnings
    if not sections["summary"] and not pre_header_lines:
        parse_warnings.append("Section 'summary' was not found or could not be parsed.")
    if not sections["experience"]:
        parse_warnings.append("Section 'experience' was not found or could not be parsed.")
    elif len(experience) == 0:
        parse_warnings.append("Experience section heading detected but no roles could be structured.")
        
    if not sections["education"]:
        parse_warnings.append("Section 'education' was not found or could not be parsed.")
    elif len(education) == 0:
        parse_warnings.append("Education section heading detected but no qualifications could be structured.")
        
    # Infer skills if explicit skills section is empty/absent
    inferred_skills = []
    if not skills:
        parse_warnings.append("No explicit skills section found; skills were inferred from experience and projects.")
        from app.utils.tech_keywords import TECH_KEYWORDS
        text_to_scan = raw_text.lower()
        for kw in TECH_KEYWORDS:
            # Build smart regex matching for each keyword
            escaped_kw = re.escape(kw.lower())
            if kw.lower() == "c":
                pattern = rf'\bc\b(?!\+|#)'
            elif kw.lower() == "net":
                pattern = rf'(?<!\.)\bnet\b'
            elif kw.isalnum():
                pattern = rf'\b{escaped_kw}\b'
            else:
                prefix = r'\b' if kw[0].isalnum() else r'(?:^|[\s\.,;\!\?])'
                suffix = r'\b' if kw[-1].isalnum() else r'(?:$|[\s\.,;\!\?])'
                pattern = rf'{prefix}{escaped_kw}{suffix}'
                
            if re.search(pattern, text_to_scan):
                inferred_skills.append(kw)

    # Parse candidate name and contact details from pre_header_lines
    candidate_name = ""
    contact_info = []
    
    # 1. Candidate's Name heuristic
    for line in pre_header_lines:
        line_str = line.strip()
        if not line_str:
            continue
        has_email = "@" in line_str
        has_phone = re.search(r'\+?\d[\d\-\s\(\)]{8,}\d', line_str) is not None
        has_url = any(u in line_str.lower() for u in ["github.com", "linkedin.com", "www."])
        
        if not (has_email or has_phone or has_url):
            if re.search(r'[a-zA-Z]', line_str):
                candidate_name = line_str
                break
                
    # 2. Contact details heuristic
    for line in pre_header_lines:
        line_str = line.strip()
        if not line_str:
            continue
        has_email = "@" in line_str
        has_phone = re.search(r'\+?\d[\d\-\s\(\)]{8,}\d', line_str) is not None
        has_url = any(u in line_str.lower() for u in ["github.com", "linkedin.com", "www."])
        
        if has_email or has_phone or has_url:
            parts = re.split(r'[\|•·,;]', line_str)
            for p in parts:
                p_clean = p.strip()
                if p_clean and p_clean not in contact_info:
                    contact_info.append(p_clean)

    return ResumeDocument(
        name=candidate_name,
        contact_info=contact_info,
        summary=summary_text,
        skills=skills,
        experience=experience,
        projects=projects,
        education=education,
        certifications=certifications,
        inferred_skills=inferred_skills,
        raw_text=raw_text,
        parse_warnings=parse_warnings
    )
